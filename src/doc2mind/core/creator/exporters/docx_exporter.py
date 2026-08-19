"""Word (DOCX) 原生文档导出引擎 — 基于 python-docx。

特性：
- 严格遵循企业级公文与深度研报排版规范（标题层级、字体行距、边距）；
- 自动识别 Markdown 标题、段落、列表项与结构化表格；
- 表格自动设置表头样式、列宽与细边框；
- 完美支持导出研报、立项方案、技术白皮书与课程教案。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from doc2mind.core.creator.models import ArtifactModel

logger = logging.getLogger("doc2mind.creator.docx")


class DocxExporter:
    """DOCX 研报/公文/教案生成器。"""

    def __init__(self) -> None:
        pass

    def export(self, artifact: ArtifactModel, output_path: Path) -> Path:
        """编译生成 .docx 物理文件并保存到指定路径。"""
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()

        # 设置页面边距（标准 A4 页边距：上2.54cm, 下2.54cm, 左3.18cm, 右3.18cm）
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

        # 调色盘
        COLOR_PRIMARY = RGBColor(30, 58, 138)  # 科技深蓝
        COLOR_TEXT = RGBColor(30, 41, 59)      # 正文深灰
        COLOR_MUTED = RGBColor(100, 116, 139)  # 注释灰

        content = artifact.raw_content
        lines = content.splitlines()

        # 主文档大标题
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(12)
        title_p.paragraph_format.space_after = Pt(8)
        run = title_p.add_run(artifact.title or "DocMind 知识研报")
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

        # 副标题/元信息
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.paragraph_format.space_after = Pt(24)
        m_run = meta_p.add_run("由 DocMind 智能知识库创作出品 · 严谨可靠")
        m_run.font.name = "Microsoft YaHei"
        m_run.font.size = Pt(10.5)
        m_run.font.color.rgb = COLOR_MUTED

        in_table = False
        table_lines: list[str] = []

        def flush_table() -> None:
            nonlocal in_table, table_lines
            if not table_lines:
                in_table = False
                return
            parsed_rows: list[list[str]] = []
            for t_line in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", t_line):
                    continue
                cols = [c.strip() for c in t_line.strip("|").split("|")]
                if any(cols):
                    parsed_rows.append(cols)

            if parsed_rows:
                num_rows = len(parsed_rows)
                num_cols = max(len(r) for r in parsed_rows)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                for r_idx, row_data in enumerate(parsed_rows):
                    row = table.rows[r_idx]
                    for c_idx, val in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = row.cells[c_idx]
                            cell.text = val
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                            # 表头着色
                            if r_idx == 0:
                                shading = parse_xml(r'<w:shd {} w:fill="E2E8F0"/>'.format(nsdecls("w")))
                                cell._tc.get_or_add_tcPr().append(shading)

                            for cp in cell.paragraphs:
                                cp.paragraph_format.space_before = Pt(4)
                                cp.paragraph_format.space_after = Pt(4)
                                for c_run in cp.runs:
                                    c_run.font.name = "Microsoft YaHei"
                                    c_run.font.size = Pt(10)
                                    if r_idx == 0:
                                        c_run.font.bold = True
                                        c_run.font.color.rgb = COLOR_PRIMARY

                doc.add_paragraph()  # 表后空行

            table_lines.clear()
            in_table = False

        for line in lines:
            ls = line.strip()

            # 表格行检测
            if ls.startswith("|") and ls.endswith("|"):
                in_table = True
                table_lines.append(ls)
                continue
            elif in_table:
                flush_table()

            if not ls:
                continue

            # 标题检测
            if ls.startswith("# "):
                h_text = ls[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(h_text)
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(16)
                r.font.bold = True
                r.font.color.rgb = COLOR_PRIMARY
            elif ls.startswith("## "):
                h_text = ls[3:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(h_text)
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(13.5)
                r.font.bold = True
                r.font.color.rgb = COLOR_PRIMARY
            elif ls.startswith("### "):
                h_text = ls[4:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(h_text)
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(12)
                r.font.bold = True
                r.font.color.rgb = COLOR_TEXT
            elif ls.startswith(("- ", "* ", "+ ", "• ")):
                bullet_text = ls[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.25
                r = p.add_run(bullet_text)
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(11)
                r.font.color.rgb = COLOR_TEXT
            elif ls.startswith("> "):
                # 引用块
                quote_text = ls[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(f"💡 {quote_text}")
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(10.5)
                r.font.italic = True
                r.font.color.rgb = COLOR_MUTED
            elif ls.startswith("---"):
                # 分页线或分隔线
                continue
            else:
                # 普通段落
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.25
                # 首行缩进两个中文字符 (约 22pt)
                p.paragraph_format.first_line_indent = Pt(22)
                r = p.add_run(ls)
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(11)
                r.font.color.rgb = COLOR_TEXT

        if in_table:
            flush_table()

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("DOCX 导出成功: %s", output_path)
        return output_path
