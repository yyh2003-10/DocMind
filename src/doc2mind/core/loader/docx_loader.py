"""Word (.docx) 加载器 — 基于 `python-docx`。

关键：必须遍历 `doc.element.body` 的 XML 子节点，按 `w:p` / `w:tbl`
交错顺序提取，才能保留正文中段落的真实顺序。
"""

from __future__ import annotations

import hashlib
from pathlib import Path

from doc2mind.core.loader.base import Loader, LoaderError
from doc2mind.core.models import (
    DocFormat,
    DocumentElement,
    ElementType,
    LoadedDocument,
)


def _heading_level(style_name: str | None) -> int | None:
    """从段落样式名提取标题层级。

    python-docx 默认样式名为 "Heading 1" / "Heading 2" / ...
    """
    if not style_name:
        return None
    name = style_name.lower()
    if not name.startswith("heading"):
        return None
    try:
        return int(name.split()[-1])
    except (ValueError, IndexError):
        return None


def _table_to_markdown(table) -> str:
    """把 python-docx `Table` 转成 Markdown 表格文本。

    Args:
        table: `docx.table.Table` 实例

    Returns:
        GitHub-flavored Markdown 表格字符串。
    """
    rows = table.rows
    if not rows:
        return ""
    lines: list[str] = []
    for i, row in enumerate(rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + "---|" * len(cells))
    return "\n".join(lines)


class DocxLoader(Loader):
    """DOCX 文档加载器（python-docx 实现，保留段落/表格交错顺序）。"""

    supported_extensions = ("docx", "doc")

    def extract(self, path: Path) -> LoadedDocument:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError as e:
            raise LoaderError(
                "python-docx 未安装。请运行：pip install python-docx"
            ) from e

        if not path.exists():
            raise LoaderError(f"文件不存在: {path}")

        try:
            data = path.read_bytes()
            file_hash = hashlib.md5(data).hexdigest()
            doc = Document(path)

            # 建立 XML 节点 → python-docx 对象的反向索引
            # paragraphs 与 tables 的 _element 即底层 w:p / w:tbl
            para_by_elem = {p._element: p for p in doc.paragraphs}
            table_by_elem = {t._element: t for t in doc.tables}

            elements: list[DocumentElement] = []
            for child in doc.element.body:
                tag = child.tag
                if tag == qn("w:p"):
                    para = para_by_elem.get(child)
                    if para is None:
                        continue
                    text = (para.text or "").strip()
                    if not text:
                        continue
                    level = _heading_level(para.style.name if para.style else None)
                    if level is not None:
                        elem_type = ElementType.HEADING
                        metadata = {
                            "type": "heading",
                            "level": level,
                            "source_format": DocFormat.DOCX.value,
                        }
                    else:
                        elem_type = ElementType.PARAGRAPH
                        metadata = {
                            "type": "paragraph",
                            "style": para.style.name if para.style else None,
                            "source_format": DocFormat.DOCX.value,
                        }
                    elements.append(
                        DocumentElement(content=text, type=elem_type, metadata=metadata)
                    )
                elif tag == qn("w:tbl"):
                    table = table_by_elem.get(child)
                    if table is None:
                        continue
                    md = _table_to_markdown(table)
                    if not md:
                        continue
                    rows = len(table.rows)
                    cols = len(table.columns) if table.columns else 0
                    elements.append(
                        DocumentElement(
                            content=md,
                            type=ElementType.TABLE,
                            metadata={
                                "type": "table",
                                "rows": rows,
                                "cols": cols,
                                "source_format": DocFormat.DOCX.value,
                            },
                        )
                    )

            return LoadedDocument(
                source=path.name,
                format=DocFormat.DOCX,
                elements=elements,
                page_count=None,  # docx 不原生分页
                size_bytes=len(data),
                file_hash=file_hash,
            )
        except LoaderError:
            raise
        except Exception as e:  # noqa: BLE001
            raise LoaderError(f"DOCX 解析失败 ({path.name}): {e}") from e
